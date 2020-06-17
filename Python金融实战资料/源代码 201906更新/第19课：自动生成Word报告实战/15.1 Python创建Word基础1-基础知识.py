'''Python商业爬虫案例实战第15讲：自动生成Word报告实战 by 王宇韬'''
#如果下面的内容被我注释掉了，大家可以选中，然后ctrl+/（Spyder中的快捷键是ctrl+1）取消注释

'''15.1节：Python创建Word基础 - 基础知识1'''

'''1.python-dox初了解'''
import docx
# 创建内存中的word文档对象
file = docx.Document()

# 写入若干段落
file.add_paragraph('螃蟹在剥我的壳，笔记本在写我')
file.add_paragraph('漫天的我落在枫叶上雪花上')
file.add_paragraph('而你在想我')

# 保存，得首先创建出保存文件夹
file.save('D:\\我的文档\\三行情书.docx')
print('Word生成完毕！')


'''2.python-docx库的基础知识'''
# 1.创建Word文档
import docx
file = docx.Document()

# 2.添加标题
file.add_heading('三行情书2', level=0)

# 3.添加段落文字
file.add_paragraph('我喜欢你')
file.add_paragraph('上一句话是假的')
file.add_paragraph('上一句话也是假的')

# 4.添加图片
file.add_picture('D:\\我的文档\\三行情书.jpg')

# 5.添加分页符
file.add_page_break()

# 6.添加表格
table = file.add_table(rows=1, cols=3)
table.cell(0,0).text = '克制'  # 第一行第一列
table.cell(0,1).text = '再克制'  # 第一行第二列
table.cell(0,2).text = '"在吗"'  # 第一行第三列

# 7.文档保存
file.save('D:\\我的文档\\三行情书2.docx')
print('三行情书2生成完毕')

# 8.读取Word文档
file = docx.Document('D:\\我的文档\\三行情书.docx')  # 打开文件demo.docx
content = []
for paragraph in file.paragraphs:
    print(paragraph.text)  # 打印各段落内容文本
    content.append(paragraph.text)

content = ' '.join(content)  # 其中单引号里的空格表示利用空格进行连接
print(content)


