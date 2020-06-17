'''Python商业爬虫案例实战第15讲：自动生成Word报告实战 by 王宇韬'''
#如果下面的内容被我注释掉了，大家可以选中，然后ctrl+/（Spyder中的快捷键是ctrl+1）取消注释

'''15.1节：Python创建Word基础 - 自动生成Word报告实战'''

import docx
from docx.shared import RGBColor
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time

# 创建一个空白Word对象，并设置好字体
file = docx.Document()
file.styles['Normal'].font.name = u'微软雅黑'  # 可换成word里面任意字体
file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 这边记得也得填一下

# 创建一个封面
p = file.add_paragraph()  # 创建一个段落
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中设置
p.paragraph_format.space_before = Pt(180)  # 段前距为180，这个是测试出来的
p.paragraph_format.space_after = Pt(30)  # 段后距为30，这个也是测试出来的
run = p.add_run('华小智舆情报告')  # 在段落里添加内容
font = run.font  # 设置字体
font.color.rgb = RGBColor(54, 95, 145)  # 颜色设置，这里是用RGB颜色
font.size = Pt(42)  # 字体大小设置，和word里面的字号相对应

# 在封面上添加日期
year = time.strftime("%Y")
month = time.strftime("%m")
day = time.strftime("%d")
today = year + '年' + month + '月' + day + '日'  # 构造当天日期
p = file.add_paragraph()  # 新建一个段落
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(today)  # 在段落中输入当天日期
font = run.font
font.color.rgb = RGBColor(54, 95, 145)
font.size = Pt(26)

# 添加分页符
file.add_page_break()

# 设置正文标题
p = file.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 段落文字居中设置
run = p.add_run('第一部分 阿里巴巴舆情报告')
run.font.color.rgb = RGBColor(54, 95, 145)  # 字体颜色设置
run.font.size = Pt(22)  # 字体大小设置

# 连接数据库 提取所有今天的"阿里巴巴"的新闻信息
import pymysql
db = pymysql.connect(host='localhost', port=3306, user='root', password='', database='pachong', charset='utf8')
company = '阿里巴巴'
today = time.strftime("%Y-%m-%d")  # 这边采用标准格式的日期格式

cur = db.cursor()  # 获取会话指针，用来调用SQL语句
sql = 'SELECT * FROM test WHERE company = %s AND date = %s'  # 编写SQL语句
cur.execute(sql, (company,today))  # 执行SQL语句
data = cur.fetchall()  # 提取所有数据，并赋值给data变量
print(data)
db.commit()  # 这个其实可以不写，因为没有改变表结构
cur.close()  # 关闭会话指针
db.close()  # 关闭数据库链接

# 编写正文内容之引言
num = len(data)
p = file.add_paragraph()  # 添加新段落
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
p.paragraph_format.first_line_indent = Inches(0.32)  # 这个控制首行缩进，
introduction = '本次舆情监控目标为阿里巴巴，主要爬取网站为百度新闻，共爬取当天新闻' + str(num) + '篇，具体新闻如下：'
p.add_run(introduction)

# 编写正文内容之具体新闻内容
for i in range(len(data)):
    p = file.add_paragraph()  # 添加新段落
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 设置两端对齐
    p.add_run(str(i + 1) + '. ' + data[i][1])  # 提取新闻标题


# 编写正文内容之表格添加
tb = file.add_table(rows=num + 1, cols=3, style='Light Shading Accent 1')  # 之前定义过num = len(data)
tb.cell(0, 0).text = '监控公司'
tb.cell(0, 1).text = '新闻标题'
tb.cell(0, 2).text = '新闻来源'
for i in range(num):  # 之前定义过num = len(data)
    tb.cell(i+1, 0).text = '阿里巴巴'
    tb.cell(i+1, 1).text = data[i][1]  # 提取新闻标题
    tb.cell(i+1, 2).text = data[i][4]  # 提取新闻来源

# 把Word文档保存，注意需提前创建好保存文件夹
file.save('D:\\我的文档\\华小智舆情报告.docx')
print('华小智舆情报告生成完毕')


